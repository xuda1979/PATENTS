#!/usr/bin/env python3
"""
Convert 交底书.tex and 检索报告.tex to high-quality DOCX using python-docx.
Strategy: Parse the LaTeX source and reconstruct it natively as a Word document,
with proper tables, headings, bold/italic, and Chinese fonts.
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKSPACE = Path(__file__).parent

# ─────────────────────────────────────────────────────────────────────────────
# Document styling helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_doc_styles(doc):
    """Set document-wide styles: A4, margins, default font."""
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    # Default paragraph font
    style = doc.styles['Normal']
    font  = style.font
    font.name = '宋体'
    font.size = Pt(12)
    # East-Asian font binding
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)

    pPr = style.element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')         # 1.5× line spacing (240 = single)
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def add_heading(doc, text, level=1):
    """Add a numbered section heading."""
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.name = '黑体'
    run.font.size = Pt(14) if level == 1 else Pt(13) if level == 2 else Pt(12)
    _set_east_asia_font(run, '黑体')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_center_title(doc, text, size=20, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '黑体'
    run.font.size = Pt(size)
    _set_east_asia_font(run, '黑体')
    return p


def _set_east_asia_font(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_para(doc, text='', bold_spans=None, italic=False, indent=False, alignment=None):
    """Add a paragraph, optionally with inline bold spans."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    _fill_para(p, text, italic=italic)
    return p


def _fill_para(p, text, italic=False):
    """Parse simple LaTeX inline markup and add runs to paragraph p."""
    # We process the text segment by segment
    segments = _parse_inline(text)
    for seg_text, seg_bold, seg_italic, seg_color in segments:
        run = p.add_run(seg_text)
        run.bold = seg_bold
        run.italic = italic or seg_italic
        run.font.name = '宋体'
        run.font.size = Pt(12)
        _set_east_asia_font(run, '宋体')
        if seg_color:
            run.font.color.rgb = RGBColor(*seg_color)


def _parse_inline(text):
    """
    Tokenise simple LaTeX inline markup into (text, bold, italic, color) tuples.
    Handles: \\textbf{}, \\textit{}, \\emph{}, \\textcolor{blue}{}, 
             \\text{}, \\mathtt{}, and strips remaining LaTeX commands.
    """
    result = []
    i = 0
    s = text

    def clean(t):
        # Remove remaining LaTeX commands we don't handle
        t = re.sub(r'\\[a-zA-Z]+\s*', ' ', t)
        t = re.sub(r'[{}$\\]', '', t)
        t = re.sub(r'\s+', ' ', t)
        return t

    # Expand common macros first
    s = s.replace(r'\SeqID', 'SeqID')
    s = s.replace(r'\HPC', 'C')
    s = s.replace(r'---', '——')
    s = s.replace(r'--', '–')
    s = s.replace(r'``', '"')
    s = s.replace(r"''", '"')

    pattern = re.compile(
        r'\\textbf\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        r'|\\textit\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        r'|\\emph\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        r'|\\textcolor\{[^}]+\}\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        r'|\\texttt\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        r'|\\text\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        r'|\\mathtt\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
    )

    last = 0
    for m in pattern.finditer(s):
        # Plain text before match
        before = s[last:m.start()]
        if before:
            result.append((clean(before), False, False, None))

        if m.group(1) is not None:   # textbf
            result.append((clean(m.group(1)), True, False, None))
        elif m.group(2) is not None: # textit
            result.append((clean(m.group(2)), False, True, None))
        elif m.group(3) is not None: # emph
            result.append((clean(m.group(3)), False, True, None))
        elif m.group(4) is not None: # textcolor blue
            result.append((clean(m.group(4)), False, False, (0, 0, 180)))
        elif m.group(5) is not None: # texttt
            result.append((clean(m.group(5)), False, False, None))
        elif m.group(6) is not None: # text{}
            result.append((clean(m.group(6)), False, False, None))
        elif m.group(7) is not None: # mathtt{}
            result.append((clean(m.group(7)), False, False, None))

        last = m.end()

    tail = s[last:]
    if tail:
        result.append((clean(tail), False, False, None))

    return result


def set_table_borders(table):
    """Apply thin borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_table_row(table, cells_text, bold_first=False, header=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells_text):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.clear()
        _fill_para(p, str(cell_text))
        for run in p.runs:
            run.font.size = Pt(11)
            if header or (bold_first and i == 0):
                run.bold = True
    return row


def add_header_row(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        if i < len(row.cells):
            cell = row.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.clear()
            _fill_para(p, h)
            for run in p.runs:
                run.font.size = Pt(11)
                run.bold = True
            # Light grey background for header
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D9D9D9')
            tcPr.append(shd)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(widths_cm[i] * 567)))  # 1cm ≈ 567 twips
                tcW.set(qn('w:type'), 'dxa')
                # Remove existing tcW
                existing = tcPr.find(qn('w:tcW'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.append(tcW)


# ─────────────────────────────────────────────────────────────────────────────
# LaTeX content parser
# ─────────────────────────────────────────────────────────────────────────────

def strip_tex_command(text):
    """Remove common LaTeX formatting commands, keep inner text."""
    text = re.sub(r'\\textbf\{([^{}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^{}]*)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^{}]*)\}', r'\1', text)
    text = re.sub(r'\\textcolor\{[^}]+\}\{([^{}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^{}]*)\}', r'\1', text)
    text = re.sub(r'\\text\{([^{}]*)\}', r'\1', text)
    text = re.sub(r'\\mathtt\{([^{}]*)\}', r'\1', text)
    text = text.replace(r'\SeqID', 'SeqID')
    text = text.replace(r'\HPC', 'C')
    text = re.sub(r'\\[a-zA-Z]+\*?\s*', '', text)
    text = re.sub(r'[{}$\\]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def parse_tabular_rows(body):
    """Parse a tabular/longtable body into list of row-lists."""
    rows = []
    # Split on \\ (row separator), but not \\\\ inside cells
    raw_rows = re.split(r'\\\\', body)
    for raw in raw_rows:
        raw = raw.strip()
        # Remove \hline and \cline
        raw = re.sub(r'\\[ch]line[^{]*\{?[^}]*\}?', '', raw).strip()
        raw = re.sub(r'\\endfirsthead|\\endhead|\\endfoot|\\endlastfoot', '', raw)
        if not raw:
            continue
        cells = [c.strip() for c in raw.split('&')]
        if any(c for c in cells):
            rows.append(cells)
    return rows


def parse_table_spec(spec):
    """Parse column spec like |p{3cm}|p{12cm}| into list of widths in cm."""
    widths = []
    for m in re.finditer(r'p\{([0-9.]+)cm\}', spec):
        widths.append(float(m.group(1)))
    if not widths:
        # Count column indicators
        cols = re.findall(r'[lcrp]', spec)
        widths = [4.0] * len(cols)
    return widths


def extract_tables_and_text(tex_body):
    """
    Yield (type, content) tuples where type is one of:
      'section', 'subsection', 'subsubsection',
      'para', 'item', 'table', 'boxed', 'empty'
    """
    # Remove comments
    tex_body = re.sub(r'%[^\n]*', '', tex_body)
    # Collapse excessive blank lines
    tex_body = re.sub(r'\n{3,}', '\n\n', tex_body)

    i = 0
    tokens = []

    # Tokenise the document into meaningful blocks
    # We'll scan linearly for environments and commands

    lines = tex_body.split('\n')
    buf = []

    def flush_buf():
        text = ' '.join(buf).strip()
        buf.clear()
        if text:
            tokens.append(('para', text))

    k = 0
    while k < len(lines):
        line = lines[k].rstrip()

        # Section headings
        m = re.match(r'\\section\*?\{(.+)\}', line)
        if m:
            flush_buf()
            tokens.append(('section', m.group(1)))
            k += 1
            continue

        m = re.match(r'\\subsection\*?\{(.+)\}', line)
        if m:
            flush_buf()
            tokens.append(('subsection', m.group(1)))
            k += 1
            continue

        m = re.match(r'\\subsubsection\*?\{(.+)\}', line)
        if m:
            flush_buf()
            tokens.append(('subsubsection', m.group(1)))
            k += 1
            continue

        # tabular / longtable environment
        m = re.match(r'\s*\\begin\{(tabular|longtable)\}\{([^}]*)\}', line)
        if m:
            flush_buf()
            env_name = m.group(1)
            spec = m.group(2)
            # Collect until \end{env_name}
            env_lines = []
            k += 1
            while k < len(lines):
                if re.search(r'\\end\{' + env_name + r'\}', lines[k]):
                    break
                env_lines.append(lines[k])
                k += 1
            k += 1  # skip \end
            body = '\n'.join(env_lines)
            tokens.append(('table', spec, body))
            continue

        # table float wrapper — just collect its inner tabular
        if re.match(r'\s*\\begin\{table', line):
            k += 1
            continue
        if re.match(r'\s*\\end\{table', line):
            k += 1
            continue

        # center environment
        if re.match(r'\s*\\begin\{center\}', line):
            k += 1
            continue
        if re.match(r'\s*\\end\{center\}', line):
            k += 1
            continue

        # itemize / enumerate
        if re.match(r'\s*\\begin\{(itemize|enumerate)', line):
            flush_buf()
            k += 1
            continue
        if re.match(r'\s*\\end\{(itemize|enumerate)', line):
            k += 1
            continue

        # item
        m = re.match(r'\s*\\item(?:\[.*?\])?\s*(.*)', line)
        if m:
            # Collect continuation lines (until next \item or \end)
            item_lines = [m.group(1)]
            k += 1
            while k < len(lines):
                nxt = lines[k].rstrip()
                if re.match(r'\s*\\item', nxt) or re.match(r'\s*\\end\{', nxt):
                    break
                item_lines.append(nxt)
                k += 1
            item_text = ' '.join(item_lines).strip()
            tokens.append(('item', item_text))
            continue

        # fbox / parbox (boxed content)
        if re.match(r'.*\\noindent\\fbox', line) or re.match(r'.*\\fbox', line):
            flush_buf()
            # Collect until matching closing brace depth
            box_lines = [line]
            k += 1
            depth = line.count('{') - line.count('}')
            while k < len(lines) and depth > 0:
                box_lines.append(lines[k])
                depth += lines[k].count('{') - lines[k].count('}')
                k += 1
            box_text = '\n'.join(box_lines)
            # Extract text from parbox
            inner = re.search(r'\\parbox\{[^}]+\}\{(.+)', box_text, re.DOTALL)
            if inner:
                box_content = inner.group(1)
                # Remove trailing brace(s)
                box_content = re.sub(r'\}\s*\}\s*$', '', box_content.strip())
                tokens.append(('boxed', box_content))
            else:
                # Fallback: just strip commands
                tokens.append(('boxed', box_text))
            continue

        # tikzpicture / algorithm / figure — skip
        m = re.match(r'\s*\\begin\{(tikzpicture|algorithm|figure|algorithmic)\}', line)
        if m:
            env_skip = m.group(1)
            k += 1
            while k < len(lines):
                if re.search(r'\\end\{' + env_skip + r'\}', lines[k]):
                    k += 1
                    break
                k += 1
            tokens.append(('para', f'[{env_skip.upper()} — 见PDF版本]'))
            continue

        # equation / align / theorem environments — keep text, remove math
        m = re.match(r'\s*\\begin\{(equation|align|theorem|lemma|definition|corollary|proposition)\*?\}', line)
        if m:
            env_math = m.group(1)
            math_lines = []
            k += 1
            while k < len(lines):
                if re.search(r'\\end\{' + env_math + r'(\*)?}', lines[k]):
                    k += 1
                    break
                math_lines.append(lines[k])
                k += 1
            math_text = ' '.join(math_lines).strip()
            math_text = strip_tex_command(math_text)
            if math_text:
                tokens.append(('para', f'[公式: {math_text}]'))
            continue

        # proof environment
        if re.match(r'\s*\\begin\{proof\}', line):
            k += 1
            continue
        if re.match(r'\s*\\end\{proof\}', line):
            k += 1
            continue

        # enumerate label redefinition, renewcommand, vspace, hspace — skip
        if re.match(r'\s*\\(renewcommand|setcounter|setlength|vspace|hspace|noindent|centering|caption)', line):
            k += 1
            continue

        # maketitle / date / author — skip
        if re.match(r'\s*\\(maketitle|date|author|title)', line):
            k += 1
            continue

        # begin/end document
        if re.match(r'\s*\\(begin|end)\{document\}', line):
            k += 1
            continue

        # Empty line = paragraph break
        if not line.strip():
            flush_buf()
            k += 1
            continue

        # Otherwise accumulate text
        buf.append(line)
        k += 1

    flush_buf()
    return tokens


# ─────────────────────────────────────────────────────────────────────────────
# Document builder
# ─────────────────────────────────────────────────────────────────────────────

def build_doc_from_tokens(doc, tokens, title_text, subtitle_text=None):
    set_doc_styles(doc)

    # Title block
    add_center_title(doc, '中国移动专利申请', size=18)
    add_center_title(doc, title_text, size=16)
    if subtitle_text:
        add_center_title(doc, subtitle_text, size=13, bold=False)
    doc.add_paragraph()  # spacer

    for tok in tokens:
        kind = tok[0]

        if kind == 'section':
            add_heading(doc, strip_tex_command(tok[1]), level=1)

        elif kind == 'subsection':
            add_heading(doc, strip_tex_command(tok[1]), level=2)

        elif kind == 'subsubsection':
            add_heading(doc, strip_tex_command(tok[1]), level=3)

        elif kind == 'para':
            text = tok[1]
            # Skip purely structural leftovers
            text = re.sub(r'\\[a-zA-Z]+\*?\s*\{[^{}]*\}\s*', ' ', text)
            text = text.strip()
            if text:
                add_para(doc, text)

        elif kind == 'item':
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1)
            _fill_para(p, tok[1])

        elif kind == 'table':
            spec  = tok[1]
            body  = tok[2]
            widths = parse_table_spec(spec)
            rows   = parse_tabular_rows(body)
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            if ncols == 0:
                continue
            # Adjust widths list to ncols
            if len(widths) < ncols:
                widths += [3.0] * (ncols - len(widths))
            widths = widths[:ncols]
            # Scale widths to fit page (16cm usable)
            total = sum(widths)
            if total > 0:
                scale = 16.0 / total
                widths = [w * scale for w in widths]

            table = doc.add_table(rows=1, cols=ncols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            first_row_is_header = False
            # Detect header row (all cells have \textbf or \hline before them)
            if rows and all('\\textbf' in c or '\\textbf' in c for c in rows[0] if c):
                first_row_is_header = True

            for ri, row_cells in enumerate(rows):
                # Pad row to ncols
                while len(row_cells) < ncols:
                    row_cells.append('')
                if ri == 0:
                    # Use the pre-created first row
                    tr = table.rows[0]
                    for ci, cell_text in enumerate(row_cells[:ncols]):
                        cell = tr.cells[ci]
                        p = cell.paragraphs[0]
                        p.clear()
                        _fill_para(p, cell_text.strip())
                        if first_row_is_header:
                            for run in p.runs:
                                run.bold = True
                                run.font.size = Pt(11)
                        else:
                            for run in p.runs:
                                run.font.size = Pt(11)
                else:
                    tr = table.add_row()
                    for ci, cell_text in enumerate(row_cells[:ncols]):
                        cell = tr.cells[ci]
                        p = cell.paragraphs[0]
                        p.clear()
                        _fill_para(p, cell_text.strip())
                        for run in p.runs:
                            run.font.size = Pt(11)

            set_table_borders(table)
            set_col_widths(table, widths)
            doc.add_paragraph()  # spacer after table

        elif kind == 'boxed':
            # Render as a bordered paragraph (simulate fbox)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.3)
            p.paragraph_format.right_indent = Cm(0.3)
            # Parse the box content into sub-tokens and render inline
            box_content = tok[1]
            # Strip environment commands
            box_content = re.sub(r'\\begin\{[^}]+\}|\\end\{[^}]+\}', '', box_content)
            _fill_para(p, box_content)
            # Add a simple border via XML
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            for edge in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '6')
                el.set(qn('w:space'), '4')
                el.set(qn('w:color'), '000000')
                pBdr.append(el)
            pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# Main conversion
# ─────────────────────────────────────────────────────────────────────────────

def extract_body(tex_src):
    """Extract content between \\begin{document} and \\end{document}."""
    m = re.search(r'\\begin\{document\}(.+?)\\end\{document\}', tex_src, re.DOTALL)
    if m:
        return m.group(1)
    return tex_src


def convert(tex_path, out_path, doc_title, doc_subtitle=None):
    print(f'\n{"="*60}')
    print(f'  Converting: {tex_path.name} → {out_path.name}')
    print(f'{"="*60}')

    src = tex_path.read_text(encoding='utf-8')
    body = extract_body(src)

    tokens = extract_tables_and_text(body)
    print(f'  Parsed {len(tokens)} tokens')

    # Count tables
    n_tables = sum(1 for t in tokens if t[0] == 'table')
    print(f'  Tables found: {n_tables}')

    doc = Document()
    build_doc_from_tokens(doc, tokens, doc_title, doc_subtitle)
    doc.save(str(out_path))
    size_kb = out_path.stat().st_size / 1024
    print(f'  ✓ Saved: {out_path.name}  ({size_kb:.0f} KB)')


if __name__ == '__main__':
    files = sys.argv[1:] if len(sys.argv) > 1 else []

    targets = [
        (
            WORKSPACE / '交底书.tex',
            WORKSPACE / '交底书.docx',
            '技术交底书',
            '动态多原语密码跳频协议（DMHP）的安全通信系统及方法'
        ),
        (
            WORKSPACE / '检索报告.tex',
            WORKSPACE / '检索报告.docx',
            '检索报告',
            '动态多原语密码跳频协议（DMHP）的安全通信系统及方法'
        ),
    ]

    if files:
        targets = [t for t in targets if t[0].name in files]

    for tex_path, out_path, title, subtitle in targets:
        if tex_path.exists():
            convert(tex_path, out_path, title, subtitle)
        else:
            print(f'  SKIP: {tex_path} not found')

    print('\n完成！')
