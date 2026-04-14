"""
Convert PDF to DOCX by rendering each page as a high-resolution image
and embedding into a Word document. This preserves exact PDF appearance.
"""
import sys
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.section import WD_ORIENT
from io import BytesIO


def pdf_to_docx(pdf_path: str, docx_path: str, dpi: int = 300):
    """Convert a PDF to DOCX by rendering pages as images."""
    doc_pdf = fitz.open(pdf_path)
    doc_word = Document()

    # A4 dimensions
    A4_WIDTH_MM = 210
    A4_HEIGHT_MM = 297

    for page_idx in range(len(doc_pdf)):
        page = doc_pdf[page_idx]

        # Set up Word section for this page
        if page_idx == 0:
            section = doc_word.sections[0]
        else:
            section = doc_word.add_section()

        # Configure section to A4 with minimal margins
        section.page_width = Mm(A4_WIDTH_MM)
        section.page_height = Mm(A4_HEIGHT_MM)
        section.top_margin = Mm(0)
        section.bottom_margin = Mm(0)
        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.header_distance = Mm(0)
        section.footer_distance = Mm(0)

        # Render page at high DPI
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)

        # Save to bytes
        img_bytes = pix.tobytes("png")
        img_stream = BytesIO(img_bytes)

        # Add image to Word document, filling the full page
        paragraph = doc_word.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(0)

        run = paragraph.add_run()
        run.add_picture(img_stream, width=Mm(A4_WIDTH_MM))

        print(f"  Page {page_idx + 1}/{len(doc_pdf)} rendered")

    doc_word.save(docx_path)
    doc_pdf.close()
    print(f"  Saved: {docx_path}")


if __name__ == "__main__":
    files = [
        ("交底书.pdf", "交底书.docx"),
        ("检索报告.pdf", "检索报告.docx"),
    ]

    for pdf, docx in files:
        print(f"Converting {pdf} -> {docx} ...")
        pdf_to_docx(pdf, docx, dpi=300)
        print()

    print("Done!")
