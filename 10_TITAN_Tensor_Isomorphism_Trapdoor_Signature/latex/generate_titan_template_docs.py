import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent.parent

JIAO_TEMPLATE = ROOT / "latex" / "output" / "1_交底书_Formatted.docx"
SEARCH_TEMPLATE = ROOT / "一种xxx方法-检索报告v0.docx"

JIAO_TEX = ROOT / "latex" / "一种基于张量同构陷门的抗量子数字签名方法、系统、设备及介质-交底书.tex"
SEARCH_TEX = ROOT / "latex" / "一种基于张量同构陷门的抗量子数字签名方法、系统、设备及介质-检索报告.tex"

OUTPUTS = [
    ROOT / "专利技术交底书.docx",
    ROOT / "现有技术检索报告.docx",
    ROOT / "latex" / "output" / "专利技术交底书.docx",
    ROOT / "latex" / "output" / "现有技术检索报告.docx",
]

TITLE = "一种基于张量同构陷门的抗量子数字签名方法、系统、设备及介质"
INVENTOR = "许达"
CONTACT = "许达 (xudayj@chinamobile.com, +86-13521894156)"
SEARCHER = "许达、xudayj@chinamobile.com、+86-13521894156"
SEARCH_DATE = "2026.03.31"


def strip_comments(text: str) -> str:
    lines = []
    for line in text.splitlines():
        if line.lstrip().startswith("%"):
            continue
        lines.append(line)
    return "\n".join(lines)


def simplify_latex(text: str) -> str:
    text = strip_comments(text)
    text = text.replace("\r\n", "\n")
    text = re.sub(r"\\begin\{document\}", "", text)
    text = re.sub(r"\\end\{document\}", "", text)
    text = re.sub(r"\\begin\{figure\}.*?\\end\{figure\}", "(请参见附图)", text, flags=re.S)
    text = re.sub(r"\\begin\{table\}.*?\\end\{table\}", "(请参见表格)", text, flags=re.S)
    text = re.sub(r"\\begin\{longtable\}\{[^}]*\}", "", text)
    text = re.sub(r"\\end\{longtable\}", "", text)
    text = re.sub(r"\\begin\{tabular\}\{[^}]*\}", "", text)
    text = re.sub(r"\\end\{tabular\}", "", text)
    text = re.sub(r"\\hline", "", text)
    text = re.sub(r"\\includegraphics(?:\[[^\]]*\])?\{[^}]+\}", "(图片)", text)
    text = re.sub(r"\\caption\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\label\{[^}]+\}", "", text)
    text = re.sub(r"\\ref\{[^}]+\}", "", text)
    text = re.sub(r"\\cite\{[^}]+\}", "", text)
    text = re.sub(r"\\vspace\*?\{[^}]+\}", "", text)
    text = re.sub(r"\\hspace\*?\{[^}]+\}", "", text)
    text = re.sub(r"\\newpage", "", text)
    text = re.sub(r"\\noindent", "", text)
    text = re.sub(r"\\maketitle", "", text)
    text = re.sub(r"\\tableofcontents", "", text)
    text = re.sub(r"\\hrule", "", text)
    text = re.sub(r"\\setlength\{[^}]+\}\{[^}]+\}", "", text)
    text = re.sub(r"\\renewcommand\{[^}]+\}\{[^}]+\}", "", text)
    text = re.sub(r"\\section\*?\{([^}]*)\}", r"\n\1\n", text)
    text = re.sub(r"\\subsection\*?\{([^}]*)\}", r"\n\1\n", text)
    text = re.sub(r"\\subsubsection\*?\{([^}]*)\}", r"\n\1\n", text)
    text = text.replace("\\\\", "\n")
    text = re.sub(r"\\begin\{enumerate\}(\[[^\]]*\])?", "", text)
    text = re.sub(r"\\end\{enumerate\}", "", text)
    text = re.sub(r"\\begin\{itemize\}", "", text)
    text = re.sub(r"\\end\{itemize\}", "", text)
    text = re.sub(r"\\item\s*", "• ", text)
    text = re.sub(r"\$(.*?)\$", r"\1", text, flags=re.S)
    text = re.sub(r"\\\[(.*?)\\\]", r"\1", text, flags=re.S)
    text = re.sub(r"\\\((.*?)\\\)", r"\1", text, flags=re.S)

    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\\[A-Za-z]+(?:\*?)\{([^{}]*)\}", r"\1", text)

    text = re.sub(r"\\[A-Za-z]+(?:\*?)", "", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("~", " ")
    text = text.replace("&", " ")
    text = re.sub(r"^\s*longtable\s*$", "", text, flags=re.M)
    text = re.sub(r"^\s*document\s*$", "", text, flags=re.M)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_sections(tex: str, starred: bool = False):
    pattern = r"\\section\*\{([^}]+)\}" if starred else r"\\section\{([^}]+)\}"
    parts = re.split(pattern, tex)
    sections = []
    for idx in range(1, len(parts), 2):
        title = parts[idx].strip()
        body = parts[idx + 1]
        sections.append((title, body))
    return sections


def nonempty_lines(text: str):
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("|p") or re.fullmatch(r"\|?p[\d\.a-zA-Z|]+", stripped):
            continue
        if "p1.2cm" in stripped or "p5.2cm" in stripped or "p7.2cm" in stripped:
            continue
        lines.append(stripped)
    return lines


def set_run_font(run, font_name="仿宋_GB2312", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def format_body_paragraph(paragraph):
    paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    if paragraph.runs:
        for run in paragraph.runs:
            set_run_font(run)


def format_title_paragraph(paragraph):
    paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    if paragraph.runs:
        for run in paragraph.runs:
            set_run_font(run, font_name="黑体", size=12, bold=True)


def replace_section(doc, header_text, new_lines, next_headers):
    paragraphs = list(doc.paragraphs)
    start_idx = None
    end_idx = None

    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if start_idx is None and header_text in text:
            start_idx = i
            continue
        if start_idx is not None and i > start_idx:
            if any(text.startswith(prefix) for prefix in next_headers):
                end_idx = i
                break

    if start_idx is None:
        raise ValueError(f"Missing header: {header_text}")
    if end_idx is None:
        end_idx = len(paragraphs)

    target_paragraphs = paragraphs[start_idx + 1 : end_idx]
    next_paragraph = paragraphs[end_idx] if end_idx < len(paragraphs) else None

    for paragraph in target_paragraphs:
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)

    if next_paragraph is None:
        for line in new_lines:
            paragraph = doc.add_paragraph(line)
            if line.endswith("：") or line.startswith("（") and line.endswith("）"):
                format_title_paragraph(paragraph)
            else:
                format_body_paragraph(paragraph)
        return

    for line in new_lines:
        paragraph = next_paragraph.insert_paragraph_before(line)
        if line.endswith("：") or line.startswith("（") and line.endswith("）"):
            format_title_paragraph(paragraph)
        else:
            format_body_paragraph(paragraph)


def build_jiaodishu_sections():
    tex = JIAO_TEX.read_text(encoding="utf-8")
    sections = {title: simplify_latex(body) for title, body in parse_sections(tex)}

    title_lines = [TITLE]
    field_lines = nonempty_lines(sections["技术领域"])
    prior_lines = nonempty_lines(sections["现有技术的技术方案"])
    problem_lines = nonempty_lines(sections["现有技术的缺点及本申请提出要解决的技术问题"])
    detailed_lines = nonempty_lines(sections["本申请提出的技术方案的详细阐述"])
    key_lines = nonempty_lines(sections["本申请提出的创新点和欲保护点"])
    advantage_lines = nonempty_lines(sections["与第三节中最接近的现有技术相比，本申请提出有何技术优点"])
    info_lines = nonempty_lines(sections["其他有助于理解本申请提出的技术资料"])
    evidence_lines = nonempty_lines(sections["本申请提出的侵权证据可获得性"])

    return {
        "一、发明名称": title_lines,
        "二、技术领域": field_lines,
        "三、现有技术的技术方案": prior_lines,
        "四、现有技术的缺点": problem_lines,
        "五、本申请提案的技术方案": detailed_lines,
        "六、本申请提案的关键点": key_lines,
        "七、与第三条": advantage_lines,
        "八、其他有助于": info_lines,
        "九、本申请提案的侵权证据可获得性": evidence_lines,
    }


def build_search_sections():
    tex = SEARCH_TEX.read_text(encoding="utf-8")
    sections = {title: simplify_latex(body) for title, body in parse_sections(tex, starred=True)}

    keyword_lines = nonempty_lines(sections["一、使用的中文与外文检索关键词"])
    docs_lines = nonempty_lines(sections["二、相关专利文献"])
    analysis_lines = nonempty_lines(sections["三、分析评价"])
    conclusion_lines = nonempty_lines(sections["四、检索结论"])

    return {
        "一、发明名称": [TITLE],
        "二、使用的中文与外文检索关键词": keyword_lines,
        "三、相关专利文献": docs_lines,
        "四、分析评述": analysis_lines,
        "五、检索结论": conclusion_lines,
    }


def build_jiaodishu():
    doc = Document(str(JIAO_TEMPLATE))
    table = doc.tables[0]
    for row in table.rows:
        label = row.cells[0].text.strip()
        if "发明名称" in label:
            row.cells[1].text = TITLE
        elif "发明人" in label:
            row.cells[1].text = INVENTOR
        elif "技术联系人" in label:
            row.cells[1].text = CONTACT

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    label = row.cells[0].text.strip()
                    if cell == row.cells[0]:
                        set_run_font(run, font_name="黑体", size=12, bold=True)
                    else:
                        set_run_font(run, font_name="仿宋_GB2312", size=12)

    sections = build_jiaodishu_sections()
    replace_section(doc, "一、发明名称", sections["一、发明名称"], ["二、", "2、"])
    replace_section(doc, "二、技术领域", sections["二、技术领域"], ["三、", "3、"])
    replace_section(doc, "三、现有技术的技术方案", sections["三、现有技术的技术方案"], ["四、", "4、"])
    replace_section(doc, "四、现有技术的缺点", sections["四、现有技术的缺点"], ["五、", "5、"])
    replace_section(doc, "五、本申请提案的技术方案", sections["五、本申请提案的技术方案"], ["六、", "6、"])
    replace_section(doc, "六、本申请提案的关键点", sections["六、本申请提案的关键点"], ["七、", "7、"])
    replace_section(doc, "七、与第三条", sections["七、与第三条"], ["八、", "8、"])
    replace_section(doc, "八、其他有助于", sections["八、其他有助于"], ["九、", "9、"])
    replace_section(doc, "九、本申请提案的侵权证据可获得性", sections["九、本申请提案的侵权证据可获得性"], ["十、"])
    return doc


def build_search_report():
    doc = Document(str(SEARCH_TEMPLATE))
    table = doc.tables[0]
    field_values = {
        "发明名称": TITLE,
        "申报单位": "研究院",
        "检索人": SEARCHER,
        "检索日期": SEARCH_DATE,
    }
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label in field_values:
            row.cells[1].text = field_values[label]

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if cell == row.cells[0]:
                        set_run_font(run, font_name="黑体", size=12, bold=True)
                    else:
                        set_run_font(run, font_name="仿宋_GB2312", size=12)

    sections = build_search_sections()
    replace_section(doc, "一、发明名称", sections["一、发明名称"], ["二、"])
    replace_section(doc, "二、使用的中文与外文检索关键词", sections["二、使用的中文与外文检索关键词"], ["三、"])
    replace_section(doc, "三、相关专利文献", sections["三、相关专利文献"], ["四、"])
    replace_section(doc, "四、分析评述", sections["四、分析评述"], ["五、"])
    replace_section(doc, "五、检索结论", sections["五、检索结论"], ["六、"])
    return doc


def save_outputs(jiaodishu_doc, search_doc):
    jiaodishu_paths = [OUTPUTS[0], OUTPUTS[2]]
    search_paths = [OUTPUTS[1], OUTPUTS[3]]
    for path in jiaodishu_paths:
        path.parent.mkdir(parents=True, exist_ok=True)
        jiaodishu_doc.save(path)
    for path in search_paths:
        path.parent.mkdir(parents=True, exist_ok=True)
        search_doc.save(path)


def main():
    jiaodishu_doc = build_jiaodishu()
    search_doc = build_search_report()
    save_outputs(jiaodishu_doc, search_doc)
    print("Generated:")
    print(ROOT / "专利技术交底书.docx")
    print(ROOT / "现有技术检索报告.docx")


if __name__ == "__main__":
    main()
