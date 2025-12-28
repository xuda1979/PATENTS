import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold

def add_markdown_paragraph(doc, text):
    # Handle Headers
    if text.startswith('# '):
        p = doc.add_heading(level=1)
        run = p.add_run(text[2:])
        set_font(run, '黑体', 16, True)
        return
    elif text.startswith('## '):
        p = doc.add_heading(level=2)
        run = p.add_run(text[3:])
        set_font(run, '黑体', 14, True)
        return
    elif text.startswith('### '):
        p = doc.add_heading(level=3)
        run = p.add_run(text[4:])
        set_font(run, '黑体', 12, True)
        return
    elif text.startswith('#### '):
        p = doc.add_heading(level=4)
        run = p.add_run(text[5:])
        set_font(run, '黑体', 12, True)
        return

    # Handle Lists
    style = None
    if text.startswith('- '):
        style = 'List Bullet'
        text = text[2:]
    elif re.match(r'^\d+\.\s', text):
        style = 'List Number'
        text = re.sub(r'^\d+\.\s', '', text)

    p = doc.add_paragraph(style=style)
    
    # Handle Bold (**text**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True)
        else:
            run = p.add_run(part)
            set_font(run)

def parse_table(doc, lines, start_index):
    # Simple table parser for Markdown tables
    # Expects header, separator, and rows
    header = lines[start_index].strip('|').split('|')
    # Skip separator line (start_index + 1)
    
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        run = hdr_cells[i].paragraphs[0].add_run(h.strip())
        set_font(run, bold=True)
        
    current_index = start_index + 2
    while current_index < len(lines):
        line = lines[current_index].strip()
        if not line.startswith('|'):
            break
        
        cols = line.strip('|').split('|')
        row_cells = table.add_row().cells
        for i, col in enumerate(cols):
            if i < len(row_cells):
                run = row_cells[i].paragraphs[0].add_run(col.strip())
                set_font(run)
        current_index += 1
        
    return current_index - 1

def convert_file(doc, filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        if line.startswith('|'):
            i = parse_table(doc, lines, i)
        else:
            add_markdown_paragraph(doc, line)
        i += 1
    
    doc.add_page_break()

def main():
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    base_dir = r'c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature'
    files = [
        '专利申请书_中文.md',
        '权利要求书_中文.md',
        '摘要_中文.md'
    ]

    for filename in files:
        print(f"Processing {filename}...")
        convert_file(doc, os.path.join(base_dir, filename))

    output_path = os.path.join(base_dir, '专利申请文件_中国移动.docx')
    doc.save(output_path)
    print(f"Successfully saved to {output_path}")

if __name__ == '__main__':
    main()
