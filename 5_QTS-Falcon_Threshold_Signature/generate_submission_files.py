import os
import re
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Length

# Configure Matplotlib to use a font that looks more like LaTeX but blends with Word (Times-like)
# 'stix' is very close to Times New Roman
try:
    plt.rcParams['mathtext.fontset'] = 'stix'
    plt.rcParams['font.family'] = 'serif'
    plt.rcParams['font.serif'] = ['Times New Roman', 'Times', 'DejaVu Serif', 'serif']
except:
    # Fallback
    plt.rcParams['mathtext.fontset'] = 'cm'

# Cache for rendered latex images to speed up
latex_cache = {}

def render_latex(formula, fontsize=12):
    if formula in latex_cache:
        return latex_cache[formula]
    
    # Setup matplotlib for latex rendering
    # Increase figsize slightly to avoid cutting off
    fig = plt.figure(figsize=(0.1, 0.1), dpi=300) 
    
    # Pre-process formula to fix common mathtext issues
    formula = formula.replace(r'\mod', r'\ \mathrm{mod}\ ')
    formula = formula.replace(r'\|', r'\|') 
    
    # Add text to get its extent
    text = f"${formula}$"
    
    # Render text with increased weight if needed, but stix is usually good
    t = fig.text(0, 0, text, fontsize=fontsize, color='black')
    
    # Get bounding box
    try:
        renderer = fig.canvas.get_renderer()
        bbox = t.get_window_extent(renderer=renderer)
    except:
        bbox = t.get_window_extent()
    
    # Resizing logic
    dpi = fig.dpi
    width_in = bbox.width / dpi
    height_in = bbox.height / dpi
    
    # Add comfortable padding
    pad = 0.05
    fig.set_size_inches(width_in + pad*2, height_in + pad*2)
    
    # Clear and redraw centered
    fig.clear()
    t = fig.text(0.5, 0.5, text, fontsize=fontsize, ha='center', va='center', color='black')
    
    # Save to buffer
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=300, transparent=True, bbox_inches='tight', pad_inches=0.02)
    plt.close(fig)
    
    buf.seek(0)
    latex_cache[formula] = buf
    return buf

def setup_document(doc):
    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    
    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

def add_text(doc, text, is_bold=False, align=None, font_size=12, style=None):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
        
    if align:
        p.alignment = align
    
    # Split by latex markers $...$
    parts = re.split(r'(\$[^\$]+\$)', text)
    
    for part in parts:
        if part.startswith('$') and part.endswith('$') and len(part) > 2:
            # It's a formula
            formula = part[1:-1]
            try:
                img_stream = render_latex(formula, fontsize=font_size)
                run = p.add_run()
                # Scale height slightly larger than font size for better visibility
                run.add_picture(img_stream, height=Pt(font_size * 1.5)) 
            except Exception as e:
                print(f"Failed to render latex: {formula}, error: {e}")
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(font_size)
        else:
            # Handle bold **text**
            subparts = re.split(r'(\*\*.*?\*\*)', part)
            for subpart in subparts:
                run = p.add_run()
                content = subpart
                
                if subpart.startswith('**') and subpart.endswith('**'):
                    content = subpart[2:-2]
                    run.bold = True
                else:
                    run.bold = is_bold
                
                run.text = content
                run.font.name = '宋体'
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def parse_table(doc, lines, start_index):
    # Parse markdown table
    # Line start_index is header: | Col1 | Col2 |
    # Line start_index+1 is separator: |---|---|
    
    header_line = lines[start_index].strip()
    # Handle both | Col1 | and Col1 | Col2 format
    if header_line.startswith('|'):
        headers = [c.strip() for c in header_line.strip('|').split('|')]
    else:
        headers = [c.strip() for c in header_line.split('|')]
    
    # Check if next line is separator
    if start_index + 1 < len(lines):
        sep_line = lines[start_index + 1].strip()
        if not '---' in sep_line:
            # Not a table, just text with pipes?
            add_text(doc, header_line)
            return start_index
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = 'Table Grid'
    except:
        # If style doesn't exist, we might get an error or plain table.
        # Python-docx defaults usually include 'Table Grid'
        pass
    
    # Fill header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        if i >= len(hdr_cells): break
        h = h.replace('**', '')
        # Clear default paragraph
        p = hdr_cells[i].paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.bold = True
        run.font.name = '宋体'
        run.font.size = Pt(10.5) # Table text often smaller (5号)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
    current_index = start_index + 2 # Skip separator
    while current_index < len(lines):
        line = lines[current_index].strip()
        if not line or (line.startswith('#') and not line.startswith('|')):
            break
        
        # It's a row if it has pipes or looks like one
        if '|' in line:
            cols = [c.strip() for c in line.strip('|').split('|')]
            row_cells = table.add_row().cells
            for i, col in enumerate(cols):
                if i >= len(row_cells): break
                col = col.replace('**', '')
                p = row_cells[i].paragraphs[0]
                p.clear()
                add_text_to_paragraph(p, col, font_size=10.5)
        else:
            # End of table
            break
                
        current_index += 1
        
    return current_index - 1

def add_text_to_paragraph(p, text, is_bold=False, font_size=12):
    # Split by latex markers $...$
    parts = re.split(r'(\$[^\$]+\$)', text)
    
    for part in parts:
        if part.startswith('$') and part.endswith('$') and len(part) > 2:
            # It's a formula
            formula = part[1:-1]
            try:
                img_stream = render_latex(formula, fontsize=font_size)
                run = p.add_run()
                # Vertical alignment of inline images is tricky in Word
                # We scale it to match text height roughly (Pt to Inch/Emu conversion handled by library)
                run.add_picture(img_stream, height=Pt(font_size * 1.5)) 
            except Exception as e:
                print(f"Failed to render latex: {formula}, error: {e}")
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(font_size)
        else:
            # Handle bold **text**
            subparts = re.split(r'(\*\*.*?\*\*)', part)
            for subpart in subparts:
                run = p.add_run()
                content = subpart
                
                # Check for bold marker
                local_bold = is_bold
                if subpart.startswith('**') and subpart.endswith('**'):
                    content = subpart[2:-2]
                    local_bold = True
                
                # Clean up any other mild markdown if present
                # e.g. [text] - keep text
                # clean_content = re.sub(r'\[(.*?)\]', r'\1', content) # Might be dangerous for citations [1]
                
                run.text = content
                run.bold = local_bold
                run.font.name = '宋体'
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_text(doc, text, is_bold=False, align=None, font_size=12, style=None):
    if style:
        try:
            p = doc.add_paragraph(style=style)
        except:
            p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        
    if align:
        p.alignment = align
    
    add_text_to_paragraph(p, text, is_bold, font_size)
    return p

def generate_claims(base_dir):
    doc = Document()
    setup_document(doc)
    
    filepath = os.path.join(base_dir, '权利要求书_中文.md')
    print(f"Generating Claims from {filepath}...")
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    next_is_claim_start = False
    claim_number = ""
    
    for line in lines:
        line = line.strip()
        if not line: continue
        if line.startswith('# 权利要求书'): continue
        if line.startswith('## 独立权利要求'): continue
        
        if "权利要求" in line and (line.startswith('###') or line.startswith('##')):
            match = re.search(r'权利要求\s*(\d+)', line)
            if match:
                claim_number = match.group(1)
                next_is_claim_start = True
            continue
            
        if next_is_claim_start:
            line = f"{claim_number}. {line}"
            next_is_claim_start = False
        
        add_text(doc, line)

    doc.save(os.path.join(base_dir, '100001_权利要求书.docx'))

def generate_description(base_dir):
    doc = Document()
    setup_document(doc)
    
    filepath = os.path.join(base_dir, 'patent_application_cn.md')
    print(f"Generating Description from {filepath}...")
    
    # Use the robust parser
    # But we need to handle the title specifically
    
    add_text(doc, "基于格密码 Falcon 算法的量子安全门限签名系统及方法", is_bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=16)
    
    # We need to skip the first few lines of the file which repeat the title
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    # Filter out the title lines manually before passing to parser logic
    # Or just use parse_markdown_content with a flag, but the file structure is specific
    # Let's just iterate and use the new logic
    
    skip_next = False
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line: 
            i+=1
            continue
            
        if line.startswith('# '): 
            i+=1
            continue
            
        if line.startswith('## 发明名称'):
            skip_next = True
            i+=1
            continue
            
        if skip_next:
            skip_next = False
            i+=1
            continue
            
        # Use the logic from parse_markdown_content for the rest
        if line.startswith('|'):
            i = parse_table(doc, lines, i)
        elif line.startswith('## '):
            text = line[3:].strip()
            add_text(doc, text, is_bold=True, font_size=14)
        elif line.startswith('### '):
            text = line[4:].strip()
            add_text(doc, text, is_bold=True, font_size=12)
        elif line.startswith('- '):
            text = line[2:].strip()
            p = add_text(doc, text)
            p.style = 'List Bullet'
        elif re.match(r'^\d+\.\s', line):
            add_text(doc, line)
        else:
            add_text(doc, line)
        i += 1
            
    # Insert Performance Charts
    doc.add_page_break()
    add_text(doc, "附：性能对比数据图表", is_bold=True, font_size=14)
    
    chart1 = os.path.join(base_dir, 'Chart_SigSize.png')
    if os.path.exists(chart1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart1, width=Cm(14.0))
        add_text(doc, "图 4：签名长度对比", align=WD_ALIGN_PARAGRAPH.CENTER)
        
    chart2 = os.path.join(base_dir, 'Chart_GasCost.png')
    if os.path.exists(chart2):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart2, width=Cm(14.0))
        add_text(doc, "图 5：链上验证 Gas 消耗对比", align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(os.path.join(base_dir, '100002_说明书.docx'))

def generate_abstract(base_dir):
    doc = Document()
    setup_document(doc)
    
    filepath = os.path.join(base_dir, '摘要_中文.md')
    print(f"Generating Abstract from {filepath}...")
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    capture = False
    abstract_text = []
    
    for line in lines:
        line = line.strip()
        if "摘要" in line and line.startswith('###'):
            capture = True
            continue
        if "关键词" in line:
            capture = False
        
        if capture and line and not line.startswith('#') and not line.startswith('---'):
            abstract_text.append(line)
            
    full_text = "\n".join(abstract_text)
    add_text(doc, full_text)
    
    doc.save(os.path.join(base_dir, '100003_说明书摘要.docx'))

def generate_drawings(base_dir):
    doc = Document()
    setup_document(doc)
    
    filepath = os.path.join(base_dir, 'drawings_specification.md')
    print(f"Generating Drawings from {filepath}...")
    
    figures = []
    current_fig = None
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if line.startswith('## 图') or line.startswith('## Figure'):
            if current_fig:
                figures.append(current_fig)
            current_fig = {'title': line.replace('#', '').strip(), 'desc': []}
        elif current_fig:
            current_fig['desc'].append(line)
            
    if current_fig:
        figures.append(current_fig)
        
    for i, fig in enumerate(figures):
        p = add_text(doc, fig['title'], is_bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        img_filename = f'Figure_{i+1}.png'
        img_path = os.path.join(base_dir, img_filename)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        
        if os.path.exists(img_path):
            try:
                run.add_picture(img_path, width=Cm(16.0))
            except Exception as e:
                print(f"Error inserting image {img_path}: {e}")
                run.add_text(f"[图片插入失败: {e}]")
        else:
            run.add_text("[图片文件未找到]")
        
        desc_text = "\n".join([l for l in fig['desc'] if l and not l.startswith('|') and not l.startswith('---')])
        add_text(doc, desc_text)
        
        doc.add_page_break()
        
    doc.save(os.path.join(base_dir, '100004_说明书附图.docx'))
    
    # Abstract Drawing
    doc_abs = Document()
    setup_document(doc_abs)
    if figures:
        fig = figures[0]
        add_text(doc_abs, fig['title'], is_bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        img_filename = 'Figure_1.png'
        img_path = os.path.join(base_dir, img_filename)
        
        p = doc_abs.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        
        if os.path.exists(img_path):
            try:
                run.add_picture(img_path, width=Cm(16.0))
            except Exception as e:
                run.add_text(f"[图片插入失败: {e}]")
        else:
            run.add_text("[图片文件未找到]")
    
    doc_abs.save(os.path.join(base_dir, '100005_摘要附图.docx'))

def generate_request(base_dir):
    doc = Document()
    setup_document(doc)
    
    filepath = os.path.join(base_dir, '专利申请书_中文.md')
    print(f"Generating Request from {filepath}...")
    
    parse_markdown_content(doc, filepath, skip_title=False)
    
    doc.save(os.path.join(base_dir, '100006_专利请求书.docx'))

def main():
    base_dir = r'c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature'
    
    generate_claims(base_dir)
    generate_description(base_dir)
    generate_abstract(base_dir)
    generate_drawings(base_dir)
    generate_request(base_dir)
    
    print("All files generated successfully.")



def parse_markdown_content(doc, filepath, skip_title=True):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Skip separators
        if line.startswith('---'):
            i += 1
            continue

        # Skip main title if requested
        if skip_title and line.startswith('# '):
            skip_title = False
            i += 1
            continue
        
        # Detect Images by Figure Header
        figure_match = re.match(r'^##\s+(Figure|图)\s*(\d+)', line, re.IGNORECASE)
        if figure_match:
            # Add the header text
            text = line.replace('#', '').strip()
            add_text(doc, text, is_bold=True, font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Try to find and insert the image
            fig_num = figure_match.group(2)
            fig_path = os.path.join(os.path.dirname(filepath), f'Figure_{fig_num}.png')
            
            if os.path.exists(fig_path):
                print(f"Inserting image: {fig_path}")
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(fig_path, width=Inches(6.0))
                except Exception as e:
                    print(f"Error inserting image {fig_path}: {e}")
            else:
                print(f"Image not found: {fig_path}")
            
            i += 1
            continue

         # Detect Charts
        if "Chart" in line or "图表" in line:
            if "Gas Cost" in line or "Gas开销" in line:
                 chart_path = os.path.join(os.path.dirname(filepath), 'Chart_GasCost.png')
                 if os.path.exists(chart_path):
                     print(f"Inserting chart: {chart_path}")
                     p = doc.add_paragraph()
                     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                     run = p.add_run()
                     run.add_picture(chart_path, width=Inches(6.0))
            elif "Signature Size" in line or "签名大小" in line:
                 chart_path = os.path.join(os.path.dirname(filepath), 'Chart_SigSize.png')
                 if os.path.exists(chart_path):
                     print(f"Inserting chart: {chart_path}")
                     p = doc.add_paragraph()
                     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                     run = p.add_run()
                     run.add_picture(chart_path, width=Inches(6.0))

        # Handle Headers
        if line.startswith('# '):
            text = line[2:].strip()
            add_text(doc, text, is_bold=True, font_size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif line.startswith('## '):
            text = line[3:].strip()
            add_text(doc, text, is_bold=True, font_size=14)
        elif line.startswith('### '):
            text = line[4:].strip()
            add_text(doc, text, is_bold=True, font_size=12)
        
        # Handle Tables
        elif line.startswith('|'):
            i = parse_table(doc, lines, i)
            
        # Handle Lists
        elif line.startswith('- '):
            text = line[2:].strip()
            p = add_text(doc, text)
            p.style = 'List Bullet'
            
        # Handle Numbered Lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            add_text(doc, f"{line.split('.')[0]}. {text}")
            
        # Normal Text
        else:
            add_text(doc, line)
            
        i += 1

if __name__ == '__main__':
    main()
