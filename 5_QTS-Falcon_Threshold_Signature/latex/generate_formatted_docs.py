import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

# =================CONFIGURATION=================
# Paths
TEX_JIAODISHU = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature\latex\交底书.tex"
TEX_JIANSUO = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature\latex\检索报告.tex"

# Template Paths (The "unrelated" files serving as templates)
TEMPLATE_JIAODISHU_PATH = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature\一种量子-经典混合上行MIMO预编码方法-交底书v0.docx"
TEMPLATE_JIANSUO_PATH = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature\一种xxx方法-检索报告v0.docx"

OUTPUT_DIR = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature"  # Output to root to overwrite or sit alongside

# =================HELPERS=================

def set_chinese_font(run, font_name='SimSun', size=12, bold=False):
    """Set font to Chinese compliant font (default Songti/SimSun)"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold

def set_par_format(paragraph, space_before=0, space_after=0, line_spacing=1.5, first_line_indent=0):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        paragraph_format.first_line_indent = Pt(first_line_indent)

def parse_tex_content(file_path):
    """ Rudimentary TeX parser to extract sections and content """
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Determine sections based on standard markers
    # We will split by \section
    sections = []
    
    # Remove preamble
    if "\\begin{document}" in content:
        content = content.split("\\begin{document}")[1]
    if "\\end{document}" in content:
        content = content.split("\\end{document}")[0]
        
    # Clean up common latex commands
    content = re.sub(r'\\maketitle', '', content)
    content = re.sub(r'\\tableofcontents', '', content)
    content = re.sub(r'\\newpage', '', content)
    
    # regex for sections
    # Structure: ('Title', 'Content')
    
    # Split by section
    parts = re.split(r'\\section\*?\{([^}]+)\}', content)
    
    # parts[0] is intro text (usually empty or ignored title block)
    # parts[1] is title 1, parts[2] is content 1, parts[3] is title 2, etc.
    
    current_title = "Intro"
    current_text = parts[0]
    
    parsed_sections = []
    
    # If using parts, skip 0
    for i in range(1, len(parts), 2):
        title = parts[i].strip()
        body = parts[i+1].strip()
        parsed_sections.append({'title': title, 'body': body})
        
    return parsed_sections

def clean_tex_body(text):
    """ Remove LaTeX formatting from body text """
    # Remove comments (lines starting with %) while keeping indentation structure if relevant, but usually we just strip them
    # Note: % inside a line is a comment too, unless escaped \%
    # Simplified approach: remove lines starting with % or whitespace then %
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('%'):
            continue
        # Remove trailing comments? Risk of removing \% if we simply split by %.
        # Assuming content doesn't have inline comments for now for safety, or minimal.
        
        # Remove separating lines like "======="
        if all(c in '= -' for c in stripped) and len(stripped) > 5:
             continue
             
        cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    # Remove subsections

    text = re.sub(r'\\subsection\*?\{([^}]+)\}', r'\n[SUB]\1\n', text)
    text = re.sub(r'\\subsubsection\*?\{([^}]+)\}', r'\n[SUBSUB]\1\n', text)
    
    # Remove itemize/enumerate
    text = re.sub(r'\\begin\{enumerate\}(\[.*?\])?', '', text)
    text = re.sub(r'\\end\{enumerate\}', '', text)
    text = re.sub(r'\\begin\{itemize\}', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    text = re.sub(r'\\item\s', '  • ', text) # Replace item with bullet
    
    # Remove bold
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    
    # Remove math mode (simple)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    text = re.sub(r'\\\[(.*?)\\\]', r'\1', text, flags=re.DOTALL)
    
    # Remove figures tables
    text = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '[表格参见PDF文件]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '[图片参见PDF文件]', text, flags=re.DOTALL)
    
    # Figures with includes generally in text
    text = re.sub(r'\\includegraphics.*?\{.*?\}', '[图片]', text)
    
    return text.strip()

def clear_document_body(doc):
    """ Clear all paragraphs and tables from the document body, preserving styles/headers """
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    for t in doc.tables:
        t._element.getparent().remove(t._element)

# =================GENERATORS=================

def generate_jiaodishu():
    print(f"Loading template from {TEMPLATE_JIAODISHU_PATH}")
    try:
        doc = Document(TEMPLATE_JIAODISHU_PATH)
        # We want to clear the body but keep headers. 
        # However, clearing body is tricky in python-docx as removing elements while iterating can be buggy.
        # A safer way is to clear text or create a new doc if template is complex.
        # But user asked to "use the template". 
        # Let's try to remove all content first.
        # Note: This is aggressive. It removes the Info Table too unless we are careful.
        # The user said "unrelated patent content", so the Info Table has "MIMO" info.
        # We need to recreate the Info Table with Falcon info.
        # So nuking the body is acceptable.
        
        # Remove all paragraphs
        for element in doc.element.body:
             doc.element.body.remove(element)
             
    except Exception as e:
        print(f"Failed to load template: {e}. creating new.")
        doc = Document()
    
    # --- Cover / Header Block (Matching Reference) ---
    # Reference: Center aligned "中国移动专利申请" "技术交底书"
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中国移动专利申请")
    set_chinese_font(run, 'SimHei', 22, True) # Large, Bold, Heiti
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("技术交底书")
    set_chinese_font(run, 'SimHei', 26, True)
    set_par_format(p, space_after=20)

    # Info Table
    # 2 columns
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(13.0)
    
    # Fill Table
    data = [
        ("公司编号", ""),
        ("发明名称", "基于格密码 Falcon 算法的量子安全门限签名系统及方法"),
        ("申报单位", "研究院"),
        ("申报类型", "发明"),
        ("发明人", "许达"),
        ("技术联系人", "许达 (xudayj@chinamobile.com, +86-13521894156)"),
        ("注意事项", "1．技术联系人应为深入了解本申请提案技术方案 Тех的技术人员。\n2．请按照集团公司提供的本技术交底书模板逐项填写。")
    ]
    
    for i, (label, content) in enumerate(data):
        row = table.rows[i]
        # Label cell
        p = row.cells[0].paragraphs[0]
        run = p.add_run(label)
        set_chinese_font(run, 'SimHei', 10.5, True) # 5号字体 (10.5pt) approx
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Content cell
        p = row.cells[1].paragraphs[0]
        run = p.add_run(content)
        set_chinese_font(run, 'SimSun', 10.5, False)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph() # Spacer

    # --- Content Sections ---
    sections = parse_tex_content(TEX_JIAODISHU)
    
    # Mapping numbers to Chinese numbers for sections
    cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    
    section_count = 0
    for sec in sections:
        title = sec['title']
        
        # Skip internal latex sections if any linger
        if "申请人信息" in title or "发明人信息" in title: 
            continue # Already in table
            
        if "发明名称" in title:
           # Usually first section after table
           pass
           
        body = clean_tex_body(sec['body'])
        
        # Print Section Header
        # Format: "一、发明名称" (SimHei, 14pt/4号, Bold)
        if title == "发明名称":
            idx_str = "一"
            section_count = 1
        elif section_count > 0:
            section_count += 1
            if section_count <= len(cn_nums):
                idx_str = cn_nums[section_count-1]
            else:
                idx_str = str(section_count)
        else:
            # Fallback if parsing order is weird
             idx_str = "•"

        p = doc.add_paragraph()
        run = p.add_run(f"{idx_str}、{title}")
        set_chinese_font(run, 'SimHei', 14, True) # 4号
        set_par_format(p, space_before=12, space_after=6)
        
        # Print Body
        # Format: SimSun, 12pt/小4, 1.5 spacing, 2 char indent
        lines = body.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            
            if line.startswith("[SUB]"):
                # Subsection
                p = doc.add_paragraph()
                run = p.add_run(line.replace("[SUB]", ""))
                set_chinese_font(run, 'SimHei', 12, True) # Bold body size
                set_par_format(p, space_before=6, space_after=0)
            elif line.startswith("[SUBSUB]"):
                p = doc.add_paragraph()
                run = p.add_run(line.replace("[SUBSUB]", ""))
                set_chinese_font(run, 'SimSun', 12, True) 
                set_par_format(p, space_before=6, space_after=0)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                set_chinese_font(run, 'SimSun', 12, False)
                set_par_format(p, line_spacing=1.5, first_line_indent=24) # 2 chars approx 24pt

    out_path = os.path.join(OUTPUT_DIR, "一种基于格密码 Falcon 算法的量子安全门限签名系统及方法-交底书.docx")
    doc.save(out_path)
    print(f"Generated {out_path}")


def generate_jiansuo():
    print(f"Loading template from {TEMPLATE_JIANSUO_PATH}")
    try:
        doc = Document(TEMPLATE_JIANSUO_PATH)
        for element in doc.element.body:
             doc.element.body.remove(element)
    except Exception as e:
        print(f"Failed to load template: {e}. creating new.")
        doc = Document()
    
    # --- Header Block (Ref: Prior Art Search Report) ---
    # Center Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("现有技术检索报告")
    set_chinese_font(run, 'SimHei', 26, True) # Close to title size
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prior Art Search Report")
    set_chinese_font(run, 'Times New Roman', 18, False)
    
    doc.add_paragraph() # Spacer
    
    # Project Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("量子安全门限 Falcon 签名系统")
    set_chinese_font(run, 'SimHei', 16, True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Quantum-Secure Threshold Falcon Signature System")
    set_chinese_font(run, 'Times New Roman', 14, False)
    
    doc.add_paragraph() # Spacer
    
    # Info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中国移动通信有限公司研究院")
    set_chinese_font(run, 'SimSun', 14, False)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026年1月")
    set_chinese_font(run, 'SimSun', 12, False)
    
    doc.add_page_break()
    
    # --- Content ---
    sections = parse_tex_content(TEX_JIANSUO)
    
    for sec in sections:
        title = sec['title']
        body = clean_tex_body(sec['body'])
        
        # Section Heading
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_chinese_font(run, 'SimHei', 14, True)
        # Add background color simulation (Green used in reference)? 
        # For simplicity, focus on font structure. Reference had "A类：..."
        
        set_par_format(p, space_before=12, space_after=6)
        
        # For Search Parameters, convert text to table if possible, but parsing that is hard without structure.
        # We will dump text with nice formatting.
        
        lines = body.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            
            if line.startswith("[SUB]"):
                p = doc.add_paragraph()
                run = p.add_run(line.replace("[SUB]", ""))
                set_chinese_font(run, 'SimHei', 12, True)
                set_par_format(p, space_before=6)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                set_chinese_font(run, 'SimSun', 12, False)
                set_par_format(p, line_spacing=1.5, first_line_indent=0) # Search report usually less indented list style

    out_path = os.path.join(OUTPUT_DIR, "一种基于格密码 Falcon 算法的量子安全门限签名系统及方法-检索报告.docx")
    doc.save(out_path)
    print(f"Generated {out_path}")

if __name__ == "__main__":
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        
    generate_jiaodishu()
    generate_jiansuo()
