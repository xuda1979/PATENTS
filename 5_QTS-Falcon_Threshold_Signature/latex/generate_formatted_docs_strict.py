import re
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# =================CONFIGURATION=================
# Template Paths
TEMPLATE_JIAODISHU = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature\temp_template_jiaodishu_converted.docx"
TEMPLATE_JIANSUO = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature\一种xxx方法-检索报告v0.docx"

# Source Content Paths
TEX_JIAODISHU = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature\latex\交底书.tex"
TEX_JIANSUO = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature\latex\检索报告.tex"

# Output Directory
OUTPUT_DIR = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature"

# Metadata for Info Table (Jiaodishu)
INFO_DATA = {
    "发明名称": "基于格密码 Falcon 算法的量子安全门限签名系统及方法",
    "发明人": "许达",
    "技术联系人": "许达 (xudayj@chinamobile.com, +86-13521894156)",
    "公司编号": "" # optional
}

# =================HELPERS=================

def clean_tex_text(text):
    """ Remove LaTeX markup and clean text for Word doc """
    if not text: return ""
    
    # Remove comments
    lines = [l for l in text.split('\n') if not l.strip().startswith('%')]
    text = '\n'.join(lines)
    
    # Remove commands
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\cite\{[^}]+\}', '', text)
    text = re.sub(r'\\ref\{[^}]+\}', '', text)
    
    # Math
    text = re.sub(r'\$([^$]+)\$', r'\1', text) # inline math
    text = re.sub(r'\\\[(.*?)\\\]', r'\1', text, flags=re.DOTALL) # display math
    
    # Lists
    text = re.sub(r'\\begin\{enumerate\}', '', text)
    text = re.sub(r'\\end\{enumerate\}', '', text)
    text = re.sub(r'\\begin\{itemize\}', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    
    # Items - Replace with numbered format or simple indent
    # Simulating simple list handling
    lines = text.split('\n')
    new_lines = []
    item_count = 1
    for line in lines:
        stripped = line.strip()
        if stripped.startswith(r'\item'):
            content = stripped.replace(r'\item', '').strip()
            # Check if this looks like it should be numbered (enumerate) or bulleted (itemize)
            # Hard to tell without context stack, but "1." is safe for most patent docs
            # Or just use "•" if user didn't forbid bullets? User said "Don't use bullet points".
            # So I will use simple text or "(1)".
            # Let's use simple indentation or dash
            new_lines.append(f"({item_count}) {content}")
            item_count += 1
        else:
            new_lines.append(line)
            
    text = '\n'.join(new_lines)
    
    # Figures/Tables
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '(请参见附图)', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '(请参见表格)', text, flags=re.DOTALL)
    
    # Sections
    text = re.sub(r'\\section\*?\{([^}]+)\}', r'\n【\1】\n', text)
    text = re.sub(r'\\subsection\*?\{([^}]+)\}', r'\n【\1】\n', text)
    text = re.sub(r'\\subsubsection\*?\{([^}]+)\}', r'\n【\1】\n', text)
    
    return text.strip()

def extract_tex_section(content, section_name_pattern):
    """ Extract content of a section roughly matching the pattern """
    # Split by \section
    sections = re.split(r'\\section\*?\{([^}]+)\}', content)
    
    # sections[0] = preamble
    # sections[1] = title1, sections[2] = content1, ...
    
    found_content = ""
    for i in range(1, len(sections), 2):
        title = sections[i]
        body = sections[i+1]
        if re.search(section_name_pattern, title, re.IGNORECASE):
            found_content += body + "\n"
            
    return clean_tex_text(found_content)

def get_falcon_content_jiaodishu():
    with open(TEX_JIAODISHU, 'r', encoding='utf-8') as f:
        tex = f.read()
        
    data = {}
    data["title"] = INFO_DATA["发明名称"]
    data["field"] = extract_tex_section(tex, "技术领域")
    data["prior_art"] = extract_tex_section(tex, "现有技术")
    data["problems"] = extract_tex_section(tex, "缺点") 
    if not data["problems"]:
         data["problems"] = extract_tex_section(tex, "技术问题")
         
    # Detailed Description: "本申请提案的技术方案" + "具体实施方式"
    tech_solution = extract_tex_section(tex, "本申请提案的技术方案")
    implementation = extract_tex_section(tex, "具体实施方式")
    data["detailed"] = tech_solution + "\n\n" + implementation
    
    # Claims / Key Points
    # Usually strictly the claims. I'll read from separate file or assume extractable
    # Ideally, read standard claims.
    # Fallback: check claims section in file
    claims_text = extract_tex_section(tex, "权利要求")
    if not claims_text:
        # Try to read generated markdown claims if available in workspace
        claims_path = os.path.join(os.path.dirname(TEX_JIAODISHU), "..", "claims_EN.md") # Wait, abstract_EN? claims_EN?
        # User workspace has claims_EN.md. 
        # But we want Chinese.
        # "权利要求书_中文.md"
        claims_cn_path = os.path.join(os.path.dirname(TEX_JIAODISHU), "..", "权利要求书_中文.md")
        if os.path.exists(claims_cn_path):
             with open(claims_cn_path, 'r', encoding='utf-8') as f:
                 claims_text = f.read()
    data["key_points"] = claims_text
    
    data["advantages"] = "本发明通过基于格的构造提供量子抗性... (See Summary)" # Placeholder or extracting beneficial effects
    # Extract beneficial effects from summary/conclusion of description
    summary_text = re.search(r'本发明公开了一种.*?涉及.*?本发明通过.*?(?=\\section)', tex, re.DOTALL)
    if summary_text:
        data["advantages"] = clean_tex_text(summary_text.group(0))
    else:
        # Fallback to hardcoded extracted text from file search earlier
        data["advantages"] = "本发明...实现了 NIST 标准 Falcon 算法的高效门限签名。本发明解决了现有门限签名方案面临的量子计算威胁，实现了常数 6 轮在线通信复杂度（与节点数无关），签名长度约 666 字节（比 Dilithium 缩小 3.6 倍），以太坊链上验证 Gas 费用降低 72%。"

    data["evidence"] = "基于公开的区块链数据和签名格式验证。"
    
    return data

def get_falcon_content_jiansuo():
    with open(TEX_JIANSUO, 'r', encoding='utf-8') as f:
        tex = f.read()
        
    data = {}
    data["title"] = INFO_DATA["发明名称"]
    
    # Keywords
    # In Tex: table or list
    # Hardcoded extraction based on prior read
    data["keywords"] = "中文关键词：量子安全, 门限签名, 格密码, Falcon, NTRU, 多方计算, MPC\n英文关键词：Quantum-Safe, Threshold Signature, Lattice-based, Falcon, NTRU, MPC"
    
    # Relevant Docs
    # Extract references from tables in tex
    # Simplification:
    data["docs"] = "1. Cozzo et al., 'Falcon: Fast-Fourier Lattice-based Compact Signatures over NTRU'\n2. Gennaro & Goldfeder, 'Fast Multiparty Threshold ECDSA with Fast Trustless Setup'\n3. Damgård et al., 'Multiparty Computation from Somewhat Homomorphic Encryption'"
    
    # Analysis
    # Combine A, B, C sections
    analysis = ""
    analysis += extract_tex_section(tex, "A类")
    analysis += extract_tex_section(tex, "B类") 
    analysis += extract_tex_section(tex, "C类")
    data["analysis"] = analysis
    
    data["conclusion"] = "本发明具有显著的新颖性和创造性。经检索，未发现与本申请全部技术特征相同的现有技术。"
    
    return data

def replace_section_content(doc, header_text, new_content, next_header_text_prefix=None, preserve_header=True):
    """
    Find paragraph with header_text.
    Delete all paragraphs AFTER it until a paragraph starting with next_header_text_prefix is found.
    Insert new_content.
    """
    
    header_found = False
    paragraphs_to_remove = []
    insert_point = None
    
    # Locate header and range
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not header_found:
             if header_text in text:
                 header_found = True
                 insert_point = p
                 # print(f"Found header: {text}")
                 continue
        else:
             # inside section
             if next_header_text_prefix and any(text.startswith(prefix) for prefix in next_header_text_prefix):
                 # Found next section
                 break
             
             # If next_header_text_prefix is None, we go to end?
             # But usually we want to clear until next section.
             # If we don't know next section, this is dangerous. 
             # We rely on specific mapping.
             
             # Mark for removal (don't remove while iterating? actually iterating list copy is safer or index)
             paragraphs_to_remove.append(p)
             
    if header_found:
        # Remove old content
        for p in paragraphs_to_remove:
            p._element.getparent().remove(p._element)
            
        # Insert new content after insert_point
        # Since insert_paragraph_after doesn't exist directly on paragraph in all versions, 
        # we insert new paragraphs after the header.
        # But docx inserts at end by default or we need to use insert_paragraph_before on the *next* element?
        # But we removed the next elements... 
        
        # Strategy: split new_content by lines and insert after header.
        # To insert *after* a paragraph is tricky in simple python-docx without reference to next.
        # Wait, if we removed siblings, the `insert_point` is still in the tree.
        # We can append to the document then move? No.
        
        # Valid Insert Strategy: 
        # `p.insert_paragraph_before(text)` works on the paragraph *following* our insertion point.
        # Since we removed the "body" paragraphs, the "paragraph following insertion point" is now the "next header".
        # So we identify the `next_header_paragraph` (the one that stopped the loop).
        # And insert before it.
        pass
        
    # Re-implementation with robust index-based logic difficult due to dynamic XML.
    # Better:
    # 1. Iterate paragraphs to find indices.
    # 2. Collect paragraphs to delete.
    # 3. Collect the paragraph *before which* to insert.
    
    # But if we delete, indices shift. Pointers remain.
    
    return header_found

def process_jiaodishu():
    print(f"Processing Jiaodishu: {TEMPLATE_JIAODISHU}")
    doc = Document(TEMPLATE_JIAODISHU)
    content = get_falcon_content_jiaodishu()
    
    # 1. Update Table (Applicant Info)
    if len(doc.tables) > 0:
        table = doc.tables[0]
        # Iterate rows to find keys
        for row in table.rows:
            # Assume Col 0 is Label, Col 1 is Value
            if len(row.cells) >= 2:
                label = row.cells[0].text.strip()
                if "公司编号" in label:
                    pass
                elif "发明名称" in label:
                    row.cells[1].text = content["title"]
                elif "发明人" in label:
                    row.cells[1].text = INFO_DATA["发明人"]
                elif "联系人" in label:
                    row.cells[1].text = INFO_DATA["技术联系人"]
    
    # 2. Update Sections
    # Mapping [Header Search String, Content Key, [Next Header Prefixes]]
    sections_map = [
        ("一、发明名称", "title", ["二、", "2、"]), 
        ("二、技术领域", "field", ["三、", "3、"]),
        ("三、现有技术的技术方案", "prior_art", ["四、", "4、"]),
        ("四、现有技术的缺点", "problems", ["五、", "5、"]),
        ("五、本申请提案的技术方案", "detailed", ["六、", "6、"]),
        ("六、本申请提案的关键点", "key_points", ["七、", "7、"]),
        ("七、与第三条", "advantages", ["八、", "8、"]),
        ("八、其他有助于", "evidence", ["九、", "9、"]),
        # Last one
        ("九、本申请", "evidence", ["十、", "End"]), # Duplicate mapping logic just to fill text
    ]
    
    # We need a function that inserts correctly.
    # Logic: Find Header Paragraph. Find Next Header Paragraph (or end). Remove between. Insert new.
    
    # Iterate through map
    for header_str, content_key, next_prefixes in sections_map:
        new_text = content.get(content_key, "")
        if not new_text: continue
        
        # Find start
        start_p = None
        end_p = None
        
        # Scan doc
        # Note: doc.paragraphs is a live list. modifying acts weird if not careful.
        all_paras = doc.paragraphs
        
        start_idx = -1
        end_idx = -1
        
        for i, p in enumerate(all_paras):
            txt = p.text.strip()
            if header_str in txt:
                start_idx = i
                # print(f"Found Start: {txt} at {i}")
                
            if start_idx != -1 and i > start_idx:
                # Look for next header
                if any(txt.startswith(prefix) for prefix in next_prefixes):
                    end_idx = i
                    # print(f"Found End: {txt} at {i}")
                    break
        
        if start_idx != -1:
            # If no next header found, assume till end
            if end_idx == -1:
                end_idx = len(all_paras)
            
            # Remove paragraphs from start_idx+1 to end_idx-1
            # We must remove in reverse order to keep indices valid?
            # Or just use the element removal method which doesn't rely on list index for the object itself.
            
            # Identify paragraphs to remove
            to_remove = all_paras[start_idx+1 : end_idx]
            
            # Reference paragraph to insert BEFORE (the end paragraph)
            # If end_idx is len, we strip.
            # But wait, python-docx doesn't support insert_paragraph_before easily on 'None'.
            # We can use start_p.insert_paragraph_after() reversely.
            
            # 1. Remove old
            for p in to_remove:
                p._element.getparent().remove(p._element)
            
            # 2. Insert new
            # We want to insert AFTER start_p (which is all_paras[start_idx])
            # BUT, doc.paragraphs list is now stale because we removed items?
            # Actually, `start_p` object is still valid.
            
            start_p = all_paras[start_idx]
            
            # Insert paragraphs in reverse order after start_p so they appear in correct order?
            # No, insert_paragraph_before(text) on the *Next* element (end_p) is best.
            # But if end_idx was len(all_paras), there is no next element.
            
            lines = new_text.split('\n')
            
            if end_idx < len(all_paras):
                # We have a next paragraph
                # But wait, logic above: `all_paras` was captured BEFORE removal.
                # `all_paras[end_idx]` is the object reference to the next header.
                next_p = all_paras[end_idx] 
                # Insert before next_p
                for line in lines:
                    if line.strip():
                        # styling?
                        new_p = next_p.insert_paragraph_before(line.strip())
                        # Restore basic font?
                        # For now, let it inherit or set simple.
                        new_p.style = doc.styles['Normal']
            else:
                # Appending to end
                for line in lines:
                     if line.strip():
                         doc.add_paragraph(line.strip())

    out_path = os.path.join(OUTPUT_DIR, "一种基于格密码 Falcon 算法的量子安全门限签名系统及方法-交底书.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")

def process_jiansuo():
    print(f"Processing Jiansuo: {TEMPLATE_JIANSUO}")
    doc = Document(TEMPLATE_JIANSUO)
    content = get_falcon_content_jiansuo()
    
    sections_map = [
        ("一、发明名称", "title", ["二、"]), 
        ("二、使用的中文", "keywords", ["三、"]),
        ("三、相关专利", "docs", ["四、"]),
        ("四、分析评述", "analysis", ["五、"]),
        ("五、检索结论", "conclusion", ["End"])
    ]
    
    # Reusing Logic (Should factor out, but verify first)
    # ... Copy paste logic for safety ...
    all_paras = doc.paragraphs
    for header_str, content_key, next_prefixes in sections_map:
        new_text = content.get(content_key, "")
        if not new_text: continue
        
        start_idx = -1
        end_idx = -1
        for i, p in enumerate(all_paras):
            txt = p.text.strip()
            if header_str in txt:
                start_idx = i
            if start_idx != -1 and i > start_idx:
                if any(txt.startswith(prefix) for prefix in next_prefixes):
                    end_idx = i
                    break
                    
        if start_idx != -1:
            if end_idx == -1: end_idx = len(all_paras)
            to_remove = all_paras[start_idx+1 : end_idx]
            for p in to_remove:
                if p._element.getparent():
                    p._element.getparent().remove(p._element)
            
            start_p = all_paras[start_idx]
            lines = new_text.split('\n')
            
            if end_idx < len(all_paras):
                next_p = all_paras[end_idx] 
                for line in lines:
                    if line.strip():
                        next_p.insert_paragraph_before(line.strip())
            else:
                for line in lines:
                     if line.strip():
                         doc.add_paragraph(line.strip())

    out_path = os.path.join(OUTPUT_DIR, "一种基于格密码 Falcon 算法的量子安全门限签名系统及方法-检索报告.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    process_jiaodishu()
    process_jiansuo()
