from docx import Document
import os

path = r"c:\Users\Lenovo\patents\5_QTS-Falcon_Threshold_Signature\temp_template_jiaodishu.docm"
try:
    doc = Document(path)
    print("SUCCESS: Loaded docm file.")
    
    print("--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(row_text)
            
    print("\n--- PARAGRAPHS (First 20) ---")
    for i, p in enumerate(doc.paragraphs[:20]):
        print(f"{i}: {p.text}")
        
except Exception as e:
    print(f"ERROR: {e}")
